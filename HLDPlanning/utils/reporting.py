# ----------------------------------------------
# utils/reporting.py
# ----------------------------------------------
import os
import csv
import datetime
from typing import Iterable, Dict, List, Any, Optional

# QGIS imports are only needed when collecting layer info
from qgis.core import QgsVectorLayer, QgsWkbTypes

try:
    import xlsxwriter  # type: ignore
except Exception:  # pragma: no cover
    xlsxwriter = None  # type: ignore

# Optional: python-docx for DOCX Job Pack output
try:
    import docx  # python-docx
    from docx.shared import Pt, Inches  # noqa: F401  (Inches kept for future use)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    docx = None  # type: ignore


__all__ = [
    "collect_job_pack_data",
    "write_job_pack",
    "write_sheet",
    "write_boq_bom_xlsx",
    "write_boq_bom_csv",
]


# ---------------------------
# Helpers: BoQ/BoM loading
# ---------------------------

def _load_boq_bom_tables(results: dict) -> Dict[str, Any]:
    """Load BoQ/BOM CSVs if present and compute totals."""

    def _read_csv(path: str) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        if not path or not os.path.exists(path):
            return rows
        with open(path, newline="", encoding="utf-8") as f:
            r = csv.DictReader(f)
            for row in r:
                rows.append(row)
        return rows

    def _totals(rows: List[Dict[str, Any]]) -> Dict[str, float]:
        per_section: Dict[str, float] = {}
        gtot = 0.0
        for r in rows:
            try:
                amt = float(r.get("amount") or 0.0)
            except Exception:
                amt = 0.0
            sect = str(r.get("section", "")).strip()
            per_section[sect] = per_section.get(sect, 0.0) + amt
            gtot += amt
        return {"per_section": per_section, "grand_total": round(gtot, 2)}

    out: Dict[str, Any] = {"boq": [], "bom": [], "boq_totals": None, "bom_totals": None}
    boq_csv = results.get("boq_csv")
    bom_csv = results.get("bom_csv")

    out["boq"] = _read_csv(boq_csv)
    out["bom"] = _read_csv(bom_csv)

    out["boq_totals"] = _totals(out["boq"]) if out["boq"] else {"per_section": {}, "grand_total": 0.0}
    out["bom_totals"] = _totals(out["bom"]) if out["bom"] else {"per_section": {}, "grand_total": 0.0}
    return out


# ---------------------------
# Job Pack (DOCX or Markdown)
# ---------------------------

def collect_job_pack_data(results: dict) -> dict:
    """
    Read produced layers and assemble a compact data dict for reports.
    Adds BoQ/BOM rows and totals if CSVs exist.

    Returns:
      {
        "layers": [ {name, path, geom, count, status}, ... ],
        "pdp_rows": [], "mfg_rows": [],
        "summary": [],
        "boq": [...], "bom": [...],
        "boq_totals": {...}, "bom_totals": {...}
      }
    """
    data = {"layers": [], "pdp_rows": [], "mfg_rows": [], "summary": []}

    def _add_row(label: str, path: str) -> Optional[QgsVectorLayer]:
        v = QgsVectorLayer(path, label, "ogr")
        ok = bool(v and v.isValid())
        status = "OK" if ok else "Missing/Invalid"
        cnt = int(v.featureCount()) if ok else 0
        geom = (QgsWkbTypes.displayString(v.wkbType()) if ok else "—")
        data["layers"].append(
            {"name": label, "path": path, "geom": geom, "count": cnt, "status": status}
        )
        return v if ok else None

    # Report existence/feature counts for everything we produced
    for k, path in (results or {}).items():
        if not path:
            continue
        _add_row(str(k), str(path))

    # Attach BoQ/BOM (if CSVs written)
    pricing = _load_boq_bom_tables(results)
    data.update(pricing)

    # Human-readable short summary
    bt = pricing.get("boq_totals", {}) or {}
    bm = pricing.get("bom_totals", {}) or {}
    data["summary"].append(f"BoQ Grand Total: {bt.get('grand_total', 0.0):,.2f}")
    data["summary"].append(f"BoM Grand Total: {bm.get('grand_total', 0.0):,.2f}")

    return data


def write_job_pack(out_path: str, proj_info: dict, data: dict, feedback=None) -> str:
    """
    Write a DOCX Job Pack if python-docx is available; otherwise write Markdown.
    Sections:
      - Cover (Project, Date, Engineer)
      - Layers & Outputs (name, status, geometry, feature count, path)
      - BoQ (table excerpt + totals)
      - BoM (table excerpt + totals)
    """
    # DOCX branch
    if docx is not None:
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
        doc = docx.Document()

        # Title
        title = doc.add_paragraph("FTTH Pre-Checks Job Pack")
        title_format = title.runs[0].font
        title_format.size = Pt(20)
        title_format.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Project info
        info = doc.add_paragraph()
        info.add_run("Project: ").bold = True
        info.add_run(proj_info.get("name", "") + "\n")
        info.add_run("Date: ").bold = True
        info.add_run(proj_info.get("date", "") or datetime.date.today().isoformat())
        info.add_run("\nEngineer: ").bold = True
        info.add_run(proj_info.get("engineer", ""))
        info.add_run("\nOutput Folder: ").bold = True
        info.add_run(proj_info.get("out_dir", ""))

        doc.add_paragraph("")

        # Layers section
        doc.add_heading("Layers & Outputs", level=2)
        tbl = doc.add_table(rows=1, cols=5)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = \
            "Name", "Status", "Geometry", "Count", "Path"
        for r in data.get("layers", []):
            row = tbl.add_row().cells
            row[0].text = str(r.get("name", ""))
            row[1].text = str(r.get("status", ""))
            row[2].text = str(r.get("geom", ""))
            row[3].text = str(int(r.get("count", 0)))
            row[4].text = str(r.get("path", ""))

        # BoQ section
        if data.get("boq"):
            doc.add_paragraph("")
            doc.add_heading("Bill of Quantities (BoQ)", level=2)
            tbl2 = doc.add_table(rows=1, cols=8)
            hdr2 = tbl2.rows[0].cells
            hdr2[0].text, hdr2[1].text, hdr2[2].text, hdr2[3].text = "Section", "Code", "Item", "Unit"
            hdr2[4].text, hdr2[5].text, hdr2[6].text, hdr2[7].text = "Quantity", "Unit price", "Amount", "Notes"
            # show first ~30 rows to keep doc readable
            for r in data["boq"][:30]:
                row = tbl2.add_row().cells
                row[0].text = str(r.get("section", ""))
                row[1].text = str(r.get("item_code", ""))
                row[2].text = str(r.get("item_name", ""))
                row[3].text = str(r.get("unit", ""))
                row[4].text = str(r.get("quantity", ""))
                row[5].text = str(r.get("unit_price", ""))
                row[6].text = str(r.get("amount", ""))
                row[7].text = str(r.get("notes", "") or "")
            # totals
            bt = data.get("boq_totals", {}) or {}
            doc.add_paragraph(f"BoQ Grand Total: {bt.get('grand_total', 0.0):,.2f}")

        # BoM section
        if data.get("bom"):
            doc.add_paragraph("")
            doc.add_heading("Bill of Materials (BoM)", level=2)
            tbl3 = doc.add_table(rows=1, cols=8)
            hdr3 = tbl3.rows[0].cells
            hdr3[0].text, hdr3[1].text, hdr3[2].text, hdr3[3].text = "Section", "Code", "Item", "Unit"
            hdr3[4].text, hdr3[5].text, hdr3[6].text, hdr3[7].text = "Quantity", "Unit price", "Amount", "Notes"
            for r in data["bom"][:30]:
                row = tbl3.add_row().cells
                row[0].text = str(r.get("section", ""))
                row[1].text = str(r.get("item_code", ""))
                row[2].text = str(r.get("item_name", ""))
                row[3].text = str(r.get("unit", ""))
                row[4].text = str(r.get("quantity", ""))
                row[5].text = str(r.get("unit_price", ""))
                row[6].text = str(r.get("amount", ""))
                row[7].text = str(r.get("notes", "") or "")
            bm = data.get("bom_totals", {}) or {}
            doc.add_paragraph(f"BoM Grand Total: {bm.get('grand_total', 0.0):,.2f}")

        # Summary bullets
        if data.get("summary"):
            doc.add_paragraph("")
            doc.add_heading("Summary", level=2)
            for line in data["summary"]:
                doc.add_paragraph(str(line), style="List Bullet")

        doc.save(out_path)
        if feedback:
            feedback.pushInfo(f"Job Pack (DOCX) written: {out_path}")
        return out_path

    # Markdown fallback
    md = [
        "# FTTH Pre-Checks Job Pack",
        f"**Project:** {proj_info.get('name','')}",
        f"**Date:** {proj_info.get('date','') or datetime.date.today().isoformat()}",
        f"**Engineer:** {proj_info.get('engineer','')}",
        f"**Output Folder:** {proj_info.get('out_dir','')}",
        "",
        "## Layers & Outputs",
    ]
    for r in data.get("layers", []):
        md.append(f"- **{r.get('name','')}** — {r.get('status','')} — {r.get('geom','')} — {int(r.get('count',0))} feats")
        md.append(f"  `{r.get('path','')}`")

    if data.get("summary"):
        md.append("")
        md.append("## Summary")
        for line in data["summary"]:
            md.append(f"- {line}")

    p = os.path.splitext(out_path)[0] + ".md"
    os.makedirs(os.path.dirname(p) or ".", exist_ok=True)
    with open(p, "w", encoding="utf-8") as f:
        f.write("\n".join(md))
    if feedback:
        feedback.pushInfo(f"Job Pack (Markdown) written: {p}")
    return p


# ---------------------------
# BoQ/BoM (XLSX + CSV)
# ---------------------------

_SECTION_ORDER = {"Trenches": 0, "Ducts": 1, "Cables": 2, "Crossings": 3, "Network Elements": 4}


def _sort_rows(rows: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Stable sort: Section (predefined order), then Item name."""
    def key_fn(r: Dict[str, Any]):
        sect = r.get("section", "")
        item = r.get("item_name", "") or r.get("item", "")
        return (_SECTION_ORDER.get(sect, 99), str(item))
    return sorted(list(rows), key=key_fn)


def write_sheet(workbook, sheet_name: str, rows: Iterable[Dict], with_totals: bool = True) -> None:
    """
    Low-level sheet writer: column headers + width + section subtotals + grand total.
    Expects row keys: section, item_code, item_name, unit, quantity, unit_price, amount, notes
    """
    ws = workbook.add_worksheet(sheet_name[:31] or "Sheet1")

    # formats
    hdr = workbook.add_format({"bold": True, "align": "left", "bg_color": "#E6F2FF", "bottom": 1})
    intfmt = workbook.add_format({"num_format": "0"})
    mfmt = workbook.add_format({"num_format": "0.00"})
    money = workbook.add_format({"num_format": "#,##0.00"})
    bold = workbook.add_format({"bold": True})

    # headers
    headers = ["Section", "Code", "Item", "Unit", "Quantity", "Unit price", "Amount", "Notes"]
    for j, h in enumerate(headers):
        ws.write(0, j, h, hdr)

    # data
    rowi = 1
    last_section = None
    totals = 0.0

    for r in rows:
        section = r.get("section", "")
        if with_totals and last_section is not None and section != last_section:
            ws.write(rowi, 0, f"Subtotal — {last_section}", bold)
            ws.write(rowi, 6, totals if totals else 0, money)
            rowi += 1
            totals = 0.0
        last_section = section

        ws.write(rowi, 0, section)
        ws.write(rowi, 1, r.get("item_code", ""))
        ws.write(rowi, 2, r.get("item_name", ""))
        ws.write(rowi, 3, r.get("unit", ""))

        # quantity
        if r.get("unit") == "m":
            try:
                ws.write_number(rowi, 4, float(r.get("quantity") or 0), mfmt)
            except Exception:
                ws.write(rowi, 4, r.get("quantity"))
        else:
            try:
                ws.write_number(rowi, 4, float(r.get("quantity") or 0), intfmt)
            except Exception:
                ws.write(rowi, 4, r.get("quantity"))

        # unit price
        up = r.get("unit_price", None)
        if up is None or str(up) == "nan":
            ws.write(rowi, 5, "")
        else:
            try:
                ws.write_number(rowi, 5, float(up), money)
            except Exception:
                ws.write(rowi, 5, up)

        # amount
        amt = r.get("amount", None)
        if amt is None or str(amt) == "nan":
            ws.write(rowi, 6, "")
        else:
            try:
                ws.write_number(rowi, 6, float(amt), money)
                totals += float(amt)
            except Exception:
                ws.write(rowi, 6, amt)

        ws.write(rowi, 7, r.get("notes", "") or "")
        rowi += 1

    # last section subtotal
    if with_totals and last_section is not None:
        ws.write(rowi, 0, f"Subtotal — {last_section}", bold)
        ws.write(rowi, 6, totals if totals else 0, money)
        rowi += 1

    # grand total
    if with_totals:
        ws.write(rowi, 0, "Grand Total", bold)
        ws.write_formula(rowi, 6, f"=SUM(G2:G{rowi})", money)

    # widths
    ws.set_column(0, 0, 18)  # Section
    ws.set_column(1, 1, 12)  # Code
    ws.set_column(2, 2, 38)  # Item
    ws.set_column(3, 3, 8)   # Unit
    ws.set_column(4, 4, 12)  # Quantity
    ws.set_column(5, 6, 14)  # Unit price, Amount
    ws.set_column(7, 7, 30)  # Notes


def write_boq_bom_xlsx(out_xlsx: str, boq_rows: Iterable[Dict], bom_rows: Iterable[Dict]) -> str:
    """
    High-level helper: writes BoQ and BoM sheets into a single XLSX like your template.
    Returns the output path.
    """
    if xlsxwriter is None:
        raise RuntimeError("xlsxwriter not available; cannot write XLSX")

    os.makedirs(os.path.dirname(out_xlsx) or ".", exist_ok=True)
    wb = xlsxwriter.Workbook(out_xlsx)

    try:
        write_sheet(wb, "BoQ", _sort_rows(boq_rows), with_totals=True)
        write_sheet(wb, "BoM", _sort_rows(bom_rows), with_totals=True)
    finally:
        wb.close()

    return out_xlsx


def write_boq_bom_csv(boq_csv: str, bom_csv: str, boq_rows: Iterable[Dict], bom_rows: Iterable[Dict]) -> None:
    """
    Thin CSV helper for BoQ/BoM (diff-friendly, always good as a fallback).
    """
    os.makedirs(os.path.dirname(boq_csv) or ".", exist_ok=True)
    os.makedirs(os.path.dirname(bom_csv) or ".", exist_ok=True)

    hdr = ["section", "item_code", "item_name", "unit", "quantity", "unit_price", "amount", "notes"]

    with open(boq_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(hdr)
        for r in _sort_rows(boq_rows):
            w.writerow([
                r.get("section", ""),
                r.get("item_code", ""),
                r.get("item_name", ""),
                r.get("unit", ""),
                r.get("quantity", ""),
                r.get("unit_price", ""),
                r.get("amount", ""),
                r.get("notes", "") or "",
            ])

    with open(bom_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(hdr)
        for r in _sort_rows(bom_rows):
            w.writerow([
                r.get("section", ""),
                r.get("item_code", ""),
                r.get("item_name", ""),
                r.get("unit", ""),
                r.get("quantity", ""),
                r.get("unit_price", ""),
                r.get("amount", ""),
                r.get("notes", "") or "",
            ])
